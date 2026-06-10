"""Genera HISTORIAS_USUARIO_MS1.docx con estimación de 1 semana (40 h)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = Path(__file__).resolve().parent.parent
OUTPUT = DOCS / "HISTORIAS_USUARIO_MS1.docx"

SEMANA_HORAS = 40
SEMANA_LABEL = "1 semana (40 horas)"

HISTORIAS = [
    {
        "id": "HU1",
        "nombre": "Acceso al sistema",
        "modulo": "Seguridad",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-01",
        "como": "Usuario registrado",
        "quiero": "Iniciar sesión con credenciales o Google y cerrar sesión",
        "para": "Acceder de forma segura a las funcionalidades del sistema según mi rol",
        "descripcion": (
            "Como usuario registrado, quiero autenticarme en la plataforma "
            "(usuario/contraseña o Google) y cerrar sesión cuando termine, "
            "para proteger mis datos y operaciones."
        ),
        "proceso": [
            "Acceder a la pantalla de inicio de sesión.",
            "Ingresar usuario y contraseña, o seleccionar Continuar con Google.",
            "Validar credenciales en POST /api/v1/auth/login o /auth/google.",
            "Redirigir al panel principal según permisos del rol.",
            "Cerrar sesión eliminando el token JWT del cliente.",
        ],
        "criterios": [
            "Permite acceso con credenciales válidas y retorna JWT.",
            "Permite acceso con Google si el email está registrado.",
            "Muestra mensaje de error con credenciales inválidas.",
            "El cierre de sesión invalida la sesión activa en el frontend.",
        ],
    },
    {
        "id": "HU2",
        "nombre": "Registro de usuarios",
        "modulo": "Seguridad",
        "horas": 2,
        "prioridad": "Alta",
        "cu": "CU-02",
        "como": "Visitante",
        "quiero": "Crear una cuenta como Cliente o Vendedor",
        "para": "Acceder al sistema sin depender del administrador",
        "descripcion": (
            "Como visitante, quiero registrarme en la plataforma eligiendo rol "
            "Cliente o Vendedor, para obtener acceso inmediato al sistema."
        ),
        "proceso": [
            "Acceder a la pantalla de registro.",
            "Completar nombre, email, teléfono, contraseña y rol.",
            "Validar unicidad del email y coincidencia de contraseñas.",
            "Crear cliente automático si el rol es CLIENTE.",
            "Emitir JWT y redirigir al panel según rol.",
        ],
        "criterios": [
            "Registra usuarios con rol CLIENTE o VENDEDOR.",
            "Rechaza registro con rol ADMIN.",
            "Rechaza email duplicado.",
            "Tras registro exitoso el usuario queda autenticado.",
        ],
    },
    {
        "id": "HU3",
        "nombre": "Administración de usuarios",
        "modulo": "Seguridad",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-03",
        "como": "Administrador",
        "quiero": "Gestionar cuentas de usuario del sistema",
        "para": "Controlar quién tiene acceso y restablecer contraseñas",
        "descripcion": (
            "Como administrador, quiero listar, crear, activar/desactivar usuarios "
            "y restablecer contraseñas, para administrar el acceso al sistema."
        ),
        "proceso": [
            "Acceder al módulo de usuarios (solo ADMIN).",
            "Listar usuarios existentes.",
            "Crear usuario con username, email, contraseña y rol.",
            "Activar o desactivar usuarios.",
            "Restablecer contraseña de vendedores o clientes.",
        ],
        "criterios": [
            "Solo ADMIN puede acceder al módulo.",
            "No permite username ni email duplicados.",
            "Usuario CLIENTE vinculado a cliente existente.",
            "Restablecer contraseña con BCrypt.",
        ],
    },
    {
        "id": "HU4",
        "nombre": "Gestión de clientes",
        "modulo": "Clientes",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-04",
        "como": "Administrador o Vendedor",
        "quiero": "Registrar y mantener la cartera de clientes",
        "para": "Gestionar las relaciones comerciales de la importadora",
        "descripcion": (
            "Como admin o vendedor, quiero registrar, editar, activar/desactivar "
            "clientes y asignarlos a mi cartera."
        ),
        "proceso": [
            "Acceder al módulo de clientes.",
            "Registrar cliente con documento, contacto y tipo (VIP, REGULAR, NUEVO).",
            "Editar y activar/desactivar clientes.",
            "Vendedor asigna cliente con Asignar a mí.",
            "Admin desactiva clientes definitivamente.",
        ],
        "criterios": [
            "ADMIN y VENDEDOR pueden crear y editar.",
            "Solo ADMIN elimina vía DELETE.",
            "No permite documento ni email duplicados.",
            "Cliente inactivo no usable en nuevos pedidos.",
        ],
    },
    {
        "id": "HU5",
        "nombre": "Consulta de perfil del cliente",
        "modulo": "Clientes",
        "horas": 1,
        "prioridad": "Media",
        "cu": "CU-05",
        "como": "Cliente",
        "quiero": "Ver mi información personal y mis pedidos",
        "para": "Hacer seguimiento de mis compras",
        "descripcion": (
            "Como cliente autenticado, quiero consultar mi perfil y pedidos "
            "asociados a mi cuenta."
        ),
        "proceso": [
            "Iniciar sesión como CLIENTE.",
            "Acceder a Mi cuenta o Mis pedidos.",
            "Sistema valida acceso solo a datos propios.",
            "Mostrar perfil y listado de pedidos.",
        ],
        "criterios": [
            "Cliente solo ve su propio perfil y pedidos.",
            "Bloquea acceso a datos ajenos (403).",
            "Muestra estado, vehículo y totales.",
        ],
    },
    {
        "id": "HU6",
        "nombre": "Consulta de catálogo",
        "modulo": "Inventario",
        "horas": 2,
        "prioridad": "Alta",
        "cu": "CU-06",
        "como": "Visitante o usuario",
        "quiero": "Ver el catálogo de vehículos disponibles",
        "para": "Conocer las opciones de compra",
        "descripcion": (
            "Como visitante, quiero listar y ver detalle de vehículos "
            "sin necesidad de registrarme."
        ),
        "proceso": [
            "Acceder al catálogo público.",
            "Listar vehículos con marca, modelo, precio e imagen.",
            "Ver detalle completo de un vehículo.",
        ],
        "criterios": [
            "Catálogo accesible sin autenticación.",
            "Detalle con VIN, precio, color y estado.",
            "Retorna 404 si no existe.",
        ],
    },
    {
        "id": "HU7",
        "nombre": "Gestión de inventario",
        "modulo": "Inventario",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-07",
        "como": "Administrador o Vendedor",
        "quiero": "Registrar, actualizar y eliminar vehículos",
        "para": "Mantener actualizado el inventario",
        "descripcion": (
            "Como admin o vendedor, quiero administrar el inventario "
            "de vehículos de la importadora."
        ),
        "proceso": [
            "Acceder al módulo de inventario.",
            "Registrar vehículo con datos técnicos y precio.",
            "Actualizar vehículos existentes.",
            "Admin elimina vehículos del sistema.",
        ],
        "criterios": [
            "ADMIN y VENDEDOR crean y editan.",
            "Solo ADMIN elimina.",
            "Vehículo nuevo en estado DISPONIBLE.",
        ],
    },
    {
        "id": "HU8",
        "nombre": "Creación de pedidos",
        "modulo": "Pedidos",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-08",
        "como": "Administrador o Vendedor",
        "quiero": "Crear un pedido vinculando cliente y vehículo",
        "para": "Iniciar el proceso comercial de venta",
        "descripcion": (
            "Como vendedor, quiero registrar un pedido asociando cliente activo "
            "con vehículo disponible y calcular el total."
        ),
        "proceso": [
            "Seleccionar cliente activo y vehículo disponible.",
            "Calcular precio + impuestos (15%) + envío.",
            "Crear pedido PENDIENTE con código PED-XXX.",
            "Reservar vehículo (RESERVADO).",
            "Notificar al cliente y alertar si falta vendedor.",
        ],
        "criterios": [
            "Rechaza cliente inactivo o vehículo no disponible.",
            "Calcula desglose financiero correcto.",
            "Genera código PED-{id}.",
            "Dispara notificaciones automáticas.",
        ],
    },
    {
        "id": "HU9",
        "nombre": "Ciclo de vida del pedido",
        "modulo": "Pedidos",
        "horas": 4,
        "prioridad": "Alta",
        "cu": "CU-09",
        "como": "Administrador o Vendedor",
        "quiero": "Avanzar o cancelar un pedido",
        "para": "Controlar el proceso hasta la entrega",
        "descripcion": (
            "Como vendedor, quiero confirmar, importar, entregar o cancelar "
            "pedidos en su flujo comercial."
        ),
        "proceso": [
            "Confirmar PENDIENTE → CONFIRMADO.",
            "Iniciar importación → EN_IMPORTACION.",
            "Entregar → ENTREGADO (vehículo VENDIDO).",
            "Cancelar o cerrar con motivo.",
            "Notificar al cliente en cada cambio.",
        ],
        "criterios": [
            "Solo transiciones de estado válidas.",
            "Entrega requiere EN_IMPORTACION.",
            "Cancelar libera vehículo a DISPONIBLE.",
        ],
    },
    {
        "id": "HU10",
        "nombre": "Asignación de pedidos",
        "modulo": "Pedidos",
        "horas": 1,
        "prioridad": "Media",
        "cu": "CU-10",
        "como": "Vendedor",
        "quiero": "Tomar pedidos sin vendedor asignado",
        "para": "Hacerme responsable de la venta",
        "descripcion": (
            "Como vendedor, quiero auto-asignarme pedidos sin responsable comercial."
        ),
        "proceso": [
            "Ver pedidos sin vendedor.",
            "Seleccionar Tomar pedido.",
            "Sistema asigna vendedor logueado.",
        ],
        "criterios": [
            "Solo VENDEDOR ejecuta la acción.",
            "Idempotente si ya tiene vendedor.",
        ],
    },
    {
        "id": "HU11",
        "nombre": "Consulta de pedidos",
        "modulo": "Pedidos",
        "horas": 2,
        "prioridad": "Alta",
        "cu": "CU-11",
        "como": "Admin, Vendedor o Cliente",
        "quiero": "Consultar pedidos según mi rol",
        "para": "Monitorear operaciones de venta",
        "descripcion": (
            "Como usuario autenticado, quiero listar pedidos filtrados por mi rol."
        ),
        "proceso": [
            "ADMIN ve todos; VENDEDOR ve propios y sin asignar.",
            "CLIENTE ve solo sus pedidos.",
            "Consultar listado o detalle por ID.",
        ],
        "criterios": [
            "Filtrado correcto por rol.",
            "Cliente no accede a pedidos ajenos.",
            "Detalle con cliente, vehículo y totales.",
        ],
    },
    {
        "id": "HU12",
        "nombre": "Trazabilidad de importaciones",
        "modulo": "Importaciones",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-12",
        "como": "Administrador",
        "quiero": "Registrar y actualizar trazabilidad aduanera",
        "para": "Controlar la logística de vehículos importados",
        "descripcion": (
            "Como administrador, quiero gestionar registros de importación "
            "con datos aduaneros y logísticos."
        ),
        "proceso": [
            "Crear importación vinculada a pedido.",
            "Actualizar estado aduanero y datos logísticos.",
            "Sincronizar con pedido y vehículo.",
            "Completar al entregar pedido.",
        ],
        "criterios": [
            "Solo ADMIN gestiona importaciones.",
            "Código IMP-{id}.",
            "No duplica importación por pedido.",
        ],
    },
    {
        "id": "HU13",
        "nombre": "Facturación de pedidos",
        "modulo": "Facturación",
        "horas": 3,
        "prioridad": "Alta",
        "cu": "CU-13",
        "como": "Administrador o Vendedor",
        "quiero": "Crear, emitir y registrar pago de facturas",
        "para": "Documentar formalmente las ventas",
        "descripcion": (
            "Como vendedor, quiero generar facturas, emitirlas y registrar su pago."
        ),
        "proceso": [
            "Crear factura BORRADOR para pedido válido.",
            "Emitir → EMITIDA.",
            "Registrar pago → PAGADA.",
        ],
        "criterios": [
            "Una factura por pedido.",
            "No factura pedidos PENDIENTES o CANCELADOS.",
            "Numeración FAC-{año}-{secuencia}.",
        ],
    },
    {
        "id": "HU14",
        "nombre": "Gestión del equipo de ventas",
        "modulo": "Vendedores",
        "horas": 2,
        "prioridad": "Media",
        "cu": "CU-14",
        "como": "Administrador",
        "quiero": "Administrar vendedores y ver KPIs",
        "para": "Supervisar el desempeño comercial",
        "descripcion": (
            "Como administrador, quiero gestionar vendedores y consultar KPIs."
        ),
        "proceso": [
            "Listar y registrar vendedores.",
            "Activar/desactivar vendedores.",
            "Consultar resumen KPIs y ranking.",
        ],
        "criterios": [
            "Solo ADMIN accede.",
            "Ranking de vendedores visible.",
        ],
    },
    {
        "id": "HU15",
        "nombre": "Centro de notificaciones",
        "modulo": "Notificaciones",
        "horas": 2,
        "prioridad": "Media",
        "cu": "CU-15",
        "como": "Usuario autenticado",
        "quiero": "Ver y gestionar mis alertas",
        "para": "Estar informado de eventos relevantes",
        "descripcion": (
            "Como usuario, quiero consultar notificaciones y marcarlas como leídas."
        ),
        "proceso": [
            "Listar alertas por rol.",
            "Contar no leídas.",
            "Marcar individual o todas como leídas.",
        ],
        "criterios": [
            "Filtradas por rol del usuario.",
            "Alertas automáticas en eventos de pedidos.",
        ],
    },
    {
        "id": "HU16",
        "nombre": "Dashboard ejecutivo",
        "modulo": "Reportes",
        "horas": 2,
        "prioridad": "Media",
        "cu": "CU-16",
        "como": "Administrador",
        "quiero": "Ver métricas globales del negocio",
        "para": "Tomar decisiones estratégicas",
        "descripcion": (
            "Como administrador, quiero panel con ventas, stock, pedidos e importaciones."
        ),
        "proceso": [
            "Acceder al dashboard (solo ADMIN).",
            "Consultar ventas, pedidos, stock y ranking.",
        ],
        "criterios": [
            "Solo ADMIN accede.",
            "Ventas de pedidos ENTREGADOS.",
            "Ranking top vendedores.",
        ],
    },
    {
        "id": "HU17",
        "nombre": "Dashboard del vendedor",
        "modulo": "Reportes",
        "horas": 1,
        "prioridad": "Media",
        "cu": "CU-17",
        "como": "Vendedor",
        "quiero": "Ver resumen de mis ventas",
        "para": "Evaluar mi desempeño comercial",
        "descripcion": (
            "Como vendedor, quiero métricas filtradas a mis operaciones."
        ),
        "proceso": [
            "Acceder a Mi resumen.",
            "Ver ventas y pedidos propios.",
        ],
        "criterios": [
            "Solo VENDEDOR accede.",
            "No expone datos de otros vendedores.",
        ],
    },
]


def set_cell_shading(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_meta_table(doc: Document, hu: dict) -> None:
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("HU-Nro.", hu["id"]),
        ("Nombre corto de HU", hu["nombre"]),
        ("Módulo", hu["modulo"]),
        ("Tiempo estimado", f"{hu['horas']} horas"),
        ("Desarrollador", "[Asignar]"),
        ("CU relacionado", hu["cu"]),
    ]
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], "D9E2F3")


def add_hu(doc: Document, hu: dict) -> None:
    doc.add_heading(f"{hu['id']} — {hu['nombre']}", level=2)
    add_meta_table(doc, hu)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Como:\t").bold = True
    p.add_run(hu["como"])
    p = doc.add_paragraph()
    p.add_run("Quiero:\t").bold = True
    p.add_run(hu["quiero"])
    p = doc.add_paragraph()
    p.add_run("Para:\t").bold = True
    p.add_run(hu["para"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Descripción:").bold = True
    doc.add_paragraph(hu["descripcion"])

    p = doc.add_paragraph()
    p.add_run("Proceso/Lógica").bold = True
    for i, step in enumerate(hu["proceso"], 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    p = doc.add_paragraph()
    p.add_run("Criterios de aceptación").bold = True
    for i, c in enumerate(hu["criterios"], 1):
        doc.add_paragraph(f"{i}. {c}", style="List Number")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].text = "Prioridad"
    cells[1].text = hu["prioridad"]
    cells[2].text = f"Estimación PHU: {hu['horas']}"
    set_cell_shading(cells[0], "D9E2F3")
    set_cell_shading(cells[2], "D9E2F3")

    doc.add_paragraph()


def add_resumen(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Resumen del backlog — MS-1", level=1)
    doc.add_paragraph(
        f"Tiempo total de desarrollo del Microservicio 1: {SEMANA_LABEL}."
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(
        ["HU-Nro.", "Nombre corto", "Módulo", "Horas", "Prioridad", "CU"]
    ):
        hdr[i].text = text
        set_cell_shading(hdr[i], "4472C4")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = None
                r.font.bold = True

    total = 0
    for hu in HISTORIAS:
        row = table.add_row().cells
        row[0].text = hu["id"]
        row[1].text = hu["nombre"]
        row[2].text = hu["modulo"]
        row[3].text = str(hu["horas"])
        row[4].text = hu["prioridad"]
        row[5].text = hu["cu"]
        total += hu["horas"]

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "TOTAL MS-1"
    total_row[2].text = SEMANA_LABEL
    total_row[3].text = str(total)
    total_row[4].text = ""
    total_row[5].text = ""
    for c in total_row:
        set_cell_shading(c, "E2EFDA")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Historias de Usuario", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Importadora de Vehículos 2026 — Microservicio 1 (MS-1 Principal)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Tiempo de desarrollo: {SEMANA_LABEL}\n").bold = True
    info.add_run("Stack: Spring Boot 3.4 · Java 17 · PostgreSQL · JWT\n")
    info.add_run("Total historias de usuario: 17")

    doc.add_paragraph()

    for hu in HISTORIAS:
        add_hu(doc, hu)

    add_resumen(doc)
    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")
    print(f"Total horas: {sum(h['horas'] for h in HISTORIAS)}")


if __name__ == "__main__":
    main()
